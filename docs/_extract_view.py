import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
d = Document('docs/SIM-PRD-VIEW-001.docx')
print('=== PARAGRAPHS ===')
for p in d.paragraphs:
    if p.text.strip():
        print(p.text)
print()
for i, t in enumerate(d.tables):
    print(f"=== TABLE {i} ===")
    for row in t.rows:
        cells = [c.text.strip().replace('\n', ' / ') for c in row.cells]
        print(" | ".join(cells))
    print()
