import io
import os


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    import fitz  # PyMuPDF
    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return '\n'.join(text_parts).strip()


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return '\n'.join(paragraphs).strip()


def parse_resume(file_path: str, file_type: str) -> str:
    """Parse a resume file and return extracted text."""
    file_type = file_type.lower().lstrip('.')
    if file_type == 'pdf':
        return parse_pdf(file_path)
    elif file_type == 'docx':
        return parse_docx(file_path)
    else:
        raise ValueError(f'Unsupported file type: {file_type}')


def allowed_file(filename: str, allowed_extensions: set) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions
